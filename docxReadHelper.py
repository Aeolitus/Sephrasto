# -*- coding: utf-8 -*-
"""
Created on Sun Mar 19 21:33:46 2017

@author: Aeolitus
"""

import docx
import re
import Fertigkeiten
import Datenbank
class tCl(object):
    def __init__(self):
        self.paragraphs = []

class w(object):
    def __init__(self, t):
        self.text = t
Reps = {
        'Pra': "V:Praioskirchliche Tradition I:1",
        'Ron': "V:Rondrakirchliche Tradition I:1",
        'Brn': "V:Boronkirchliche Tradition I:1",
        'Phe': "V:Phexkirchliche Tradition I:1",
        'Eff': "V:Efferdkirchliche Tradition I:1",
        'Hes': "V:Hesindekirchliche Tradition I:1",
        'Ing': "V:Ingerimmkirchliche Tradition I:1",
        'Ang': "V:Angroschkirchliche Tradition I:1",
        'Per': "V:Perainekirchliche Tradition I:1",
        'Fir': "V:Firunkirchliche Tradition I:1",
        'Rah': "V:Rahjakirchliche Tradition I:1",
        'Tsa': "V:Tsakirchliche Tradition I:1",
        'Tra': "V:Traviakirchliche Tradition I:1",
        'Mag': "V:Gildenmagische Repräsentation I:1",
        'Hex': "V:Hexische Repräsentation I:1",
        'Geo': "V:Geodische Repräsentation I:1",
        'Dru': "V:Druidische Repräsentation I:1",
        'Srl': "V:Scharlatanische Repräsentation I:1",
        'Sch': "V:Schelmische Repräsentation I:1",
        'Ach': "V:Kristallomantische Repräsentation I:1",
        'Bor': "V:Borbaradianische Repräsentation I:1",
        'Alch': "V:Alchemistische Repräsentation I:1",
        'Elf': "V:Elfische Repräsentation I:1",
        'Swa': "V:Swafnirkirchliche Tradition I:1",
        'Ifi': "V:Ifirnkirchliche Tradition I:1",
        'Nan': "V:Nanduskirchliche Tradition I:1"
        } 
def getText(filename):
    doc = docx.Document(filename)
    fullText = tCl()
    for para in doc.paragraphs:
        for el in para.text.split('\n'):
            t = w(el)
            fullText.paragraphs.append(t)
    #return '\n'.join(fullText)
    return fullText
def updateUeber():
    ferts = {}
    txt = getText('ueber_new.docx')
    #results = re.search('', txt)
    count = 0
    while True:
        res = re.match("(?P<name>.*?)\s*?\(\s*?(?P<at1>[A-Z]{2})\s*?/\s*?(?P<at2>[A-Z]{2})\s*?/\s*?(?P<at3>[A-Z]{2})\s*?,\s*?(?P<sf>[0-9])\s*?\)\s*?", txt.paragraphs[count].text, re.UNICODE)
        if res is not None:
            fert = Fertigkeiten.Fertigkeit()
            dic = res.groupdict()
            fert.name = dic['name']
            fert.steigerungsfaktor = int(dic['sf'])
            fert.attribute = [dic['at1'], dic['at2'], dic['at3']]
            count += 1
            fert.text = txt.paragraphs[count].text
            ferts.update({dic['name']: fert})                         
        count +=1
        if count >= len(txt.paragraphs):
            break
    D = Datenbank.Datenbank()
    D.vorteile.clear()
    D.talente.clear()
    D.fertigkeiten.clear()
    D.übernatürlicheFertigkeiten.clear()
    D.datei = "datenbank2.xml"
    D.übernatürlicheFertigkeiten.update(ferts)
    D.xmlSchreiben()
    return ferts

def getMirakel():
    txt = getText('ueber_new.docx')
    #results = re.search('', txt)
    count = len(txt.paragraphs)-1
    res = []
    fer = []
    dic = {}
    lastR = ""
    taken = True
    while count > 0:
        rer = re.findall("Mirakel\s*?[(](?P<ferts>[a-zA-ZäöüÄÖÜß ,()]+)[)]",txt.paragraphs[count].text,re.UNICODE)
        if len(rer) > 0:
            for el in rer:
                if taken == False:
                    print("One forgotten at " + lastR)
                res.append(el)
                lastR = el
                taken = False
        else:
            ref = re.findall("(?P<name>.*?)\s*?\(\s*?(?:[A-Z]{2})\s*?/\s*?(?:[A-Z]{2})\s*?/\s*?(?:[A-Z]{2})\s*?,\s*?(?:[0-9])\s*?\)\s*?", txt.paragraphs[count].text, re.UNICODE)
            for el in ref:
                fer.append(el)
                if taken == True:
                    print("Double Accept at " + el)
                dic[el] = lastR
                taken = True
                if el == "Schlaf":
                    count = 0;
        count -= 1
    # Load DB2
    D = Datenbank.Datenbank()
    D.datei = "datenbank3.xml"
    D.vorteile.clear()
    D.talente.clear()
    D.fertigkeiten.clear()
    D.übernatürlicheFertigkeiten.clear()
    #D.xmlLaden()
    
    # For each entry in dict, create Mirakel Talent
    # Assign Mirakel to different Ferts as well! 
    splitDic = {}
    for el in dic:
        for el2 in dic[el].split(","):
            el3 = el2.strip()
            if el3 == "Wurfwaffen(nur Fir)":
                el3 = "Wurfwaffen (nur Fir)"
            if el3 == "Wahrnehmung (nur Ifi))":
                el3 = el3[:-1]
            if el3 == "Tiergestalt (Tod":
                el3 += ")"
            idx = el3.find(" (")
            if idx != -1:
                el3 = el3[:idx]
            if el3 in splitDic:
                tmp = splitDic[el3]
                tmp.append(el)
                splitDic[el3] = tmp
            else:    
                splitDic[el3] = [el]
    print("""Fix Manually:
    CH nur Ifi Wildnis
    Heilkunde nur Ifi Wildnis
    IN nur Ifi Jagd
    MU nur Swa Wind und Wogen
    Schusswaffen nur Fir Jagd
    Selbstbeherrschung nur Swa Wind und Wogen
    Wahrnehmung nur Eff Wind und Wogen
    Wahrnehmung nur Ifi Jagd
    Wurfwaffen nur Fir Jagd""")
    for el in splitDic:
        tal = Fertigkeiten.Talent()
        tal.name = "Mirakel: " + el
        tal.fertigkeiten = splitDic[el]
        tal.kosten = 20
        tal.text = "Dein nächster Wurf auf " + el + """ ist um +4 Punkte erleichtert.
Mächtige Liturgie: Der Wurf ist um zusätzliche +2 erleichtert.
Probenschwierigkeit: 12
Vorbereitungszeit: 0 Aktionen
Zielobjekt: selbst
Reichweite: Berührung
Wirkungsdauer: 4 Minuten
Kosten: 4 KaP
Fertigkeiten: """ + ", ".join(tal.fertigkeiten)
        D.talente[tal.name] = tal
    D.xmlSchreiben()
    #return (dic,res,fer)

def updateUeberTalents():
    tals = {}
    weitere = []
    anm = None
    txt = getText('ueber_new.docx')
    #results = re.search('', txt)
    count = len(txt.paragraphs)-1
    while True:
        flag = False
        #Is this erlernen?
        resEr = re.match(r"\s*?Erlernen:\s*?(?P<verb>.*?);\s*?(?P<EP>[0-9]{1,3})\s*?EP\s*?", txt.paragraphs[count].text, re.UNICODE)
        if resEr is not None:
            tal = Fertigkeiten.Talent()
            textpart = []
            erDict = resEr.groupdict()
            tal.kosten = erDict['EP']
            # Get out different representations
            resVe = re.findall(u"(?:[\s0-9;,]*)([a-zA-ZäöüÄÖÜß]{3,4})(?:[\s0-9;,]*)",erDict['verb'], re.UNICODE)
            arr = []
            for el in resVe:
                arr.append([Reps[el]])
            tal.voraussetzungen.append(arr)
            #Move to fertigkeiten
            count -= 1
            if anm is not None:
                textpart.append(anm)
                anm = None
            textpart.append(txt.paragraphs[count].text)
            resFe = re.findall(u"[a-zA-ZäöüÄÖÜß ]+",txt.paragraphs[count].text, re.UNICODE)
            if resFe[0] != "Fertigkeiten":
                flag = True
            else:
                for el in resFe[1:]:
                    tal.fertigkeiten.append(el.strip())
            if "Abu al" in tal.fertigkeiten and "Mada"in tal.fertigkeiten:
                tal.fertigkeiten.remove("Abu al")
                tal.fertigkeiten.remove("Mada")
                tal.fertigkeiten.insert(0,"Abu al'Mada")
            count -= 1
            #Get Text
            iterate = txt.paragraphs[count].text
            while not iterate.startswith("Erlernen:"):
                textpart.append(txt.paragraphs[count].text)
                if iterate == "Bannschwert" or iterate == "Eidsegen":
                    break
                count -= 1
                iterate = txt.paragraphs[count].text
#==============================================================================
#             while not txt.paragraphs[count].text.startswith("Probenschwierigkeit:"):
#                 textpart.append(txt.paragraphs[count].text)
#                 count -= 1
#             textpart.append(txt.paragraphs[count].text)
#             count -=1
#             textpart.append(txt.paragraphs[count].text)
#             if txt.paragraphs[count].text.startswith("Mächtige Magie:"):
#                 count -= 1
#                 textpart.append(txt.paragraphs[count].text)
#             if txt.paragraphs[count].text.startswith("Mächtige Liturgie:"):
#                 count -= 1
#                 textpart.append(txt.paragraphs[count].text)
#             if txt.paragraphs[count].text.startswith("Merkmal"):
#                 count -= 1
#                 textpart.append(txt.paragraphs[count].text)
#             count -= 1
#             textpart.append(txt.paragraphs[count].text)    
#==============================================================================
#==============================================================================
#             while re.match(r"[a-zA-ZäöüÄÖÜß ]:\s*?(.*?)", txt.paragraphs[count].text, re.UNICODE) is not None:
#                 textpart.append(txt.paragraphs[count].text)
#                 count -= 1
#             textpart.append(txt.paragraphs[count].text)
#             count -= 1
#             textpart.append(txt.paragraphs[count].text)
#==============================================================================
            tal.name = textpart[-1].strip()
            if tal.name.startswith("Anmerkung:") or tal.name.startswith("Anmerkungen:"):
                anm = tal.name
                textpart = textpart[:-1]
                tal.name = textpart[-1].strip()
            if tal.name.startswith("Weitere Talente:"):
                weitere.append(tal.name)
                textpart = textpart[:-1]
                tal.name = textpart[-1].strip()
            if re.match("(?P<name>.*?)\s*?\(\s*?(?P<at1>[A-Z]{2})\s*?/\s*?(?P<at2>[A-Z]{2})\s*?/\s*?(?P<at3>[A-Z]{2})\s*?,\s*?(?P<sf>[0-9])\s*?\)\s*?", tal.name, re.UNICODE) is not None:
                textpart = textpart[:-2]
                tal.name = textpart[-1].strip()
            textpart = textpart[:-1]
            textpart.reverse()
            tal.text = "\n".join(textpart)
            tals.update({tal.name: tal})
            if flag:
                print("Recheck " + tal.name)
            else:
                #print("Success with: " + tal.name)
                pass
        else:
            count -= 1
        if count <= 1:
            break
    D = Datenbank.Datenbank()
    D.datei = "datenbank2.xml"
    D.vorteile.clear()
    D.talente.clear()
    D.fertigkeiten.clear()
    D.übernatürlicheFertigkeiten.clear()
    D.xmlLaden()
    D.talente.clear()
    tmp = []
    for tal in tals:
        tmp.append(tal)
    tmp.reverse()
    tmp2 = {}
    for el in tmp:
        tmp2.update({el: tals[el]})
    D.talente.update(tmp2)
    D.xmlSchreiben()
    print("Passive Fähigkeiten stimmt noch nicht, braucht noch die Ferts!!")
    print("Fix -alle Elemente- bei Transmutation der Elemente")
    print("Fix Einzelelement bei Herbeirufung")
    print("Fix Borbaradianisch -> Boronkirchlich")
    print("Fix 'Entspricht dem Zauber' and stuff")
    print("Fix Oracnofaxius und Segen der heiligen Ardare Ferts")
    print("Fix Requirements for Traditionen!")
    print("Fix Adlerschwinge und Accuratum Anmerkung")
    print("Fix Verbreitungen? Igniplano etc")
    print("Dämonisch, nicht in Geo, Mag")
    '''
    , Kein Vorteil Praioskirchliche Tradition I, Kein Vorteil Boronkirchliche Tradition I, Kein Vorteil Rondrakirchliche Tradition I, Kein Vorteil Boronkirchliche Tradition I, Kein Vorteil Phexkirchliche Tradition I, Kein Vorteil Hesindekirchliche Tradition I, Kein Vorteil Ingerimmkirchliche Tradition I, Kein Vorteil Angroschkirchliche Tradition I, Kein Vorteil Perainekirchliche Tradition I, Kein Vorteil Firunkirchliche Tradition I, Kein Vorteil Rahjakirchliche Tradition I, Kein Vorteil Tsakirchliche Tradition I, Kein Vorteil Traviakirchliche Tradition I
    '''
    
    return tmp2

if __name__ == "__main__":
    pass
    #ret = updateUeberTalents()